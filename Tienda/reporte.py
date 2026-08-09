from docx import Document
from docx.shared import Pt
import funciones

def generarReporteVentas(conexionBD):
    print("\t\t....:::: GENERANDO REPORTE EN WORD ::::...\n")
    
    if conexionBD and conexionBD.is_connected():
        try:
            cursor = conexionBD.cursor()


            sql = """
            SELECT 
                v.id,
                v.album,
                v.artista,
                v.genero,
                v.anio,
                v.precio,
                v.stock
            FROM vinilos v
            """

            cursor.execute(sql)
            vinilos = cursor.fetchall()

            documento = Document()

            titulo = documento.add_heading("TIENDA DE VINILOS", level=1)
            titulo.runs[0].font.size = Pt(18)

            documento.add_heading("REPORTE DE INVENTARIO Y VENTAS", level=2)
            documento.add_paragraph()

            total_inventario = 0

            for item in vinilos:
                documento.add_paragraph(f"ID Vinilo: {item[0]}")
                documento.add_paragraph(f"Álbum: {item[1]}")
                documento.add_paragraph(f"Artista: {item[2]}")
                documento.add_paragraph(f"Género: {item[3]}")
                documento.add_paragraph(f"Año: {item[4]}")
                documento.add_paragraph(f"Precio Unitario: ${item[5]:.2f}")
                documento.add_paragraph(f"Stock Disponible: {item[6]}")
                
            
                subtotal_item=float(item[5]) * int(item[6])
                total_inventario+=subtotal_item
                
                documento.add_paragraph(f"Valor en Stock: ${subtotal_item:.2f}")
                documento.add_paragraph("--------------------------------------------")

            documento.add_heading(
                f"VALOR TOTAL DEL INVENTARIO: ${total_inventario:.2f}", 
                level=2
            )


            nombre_archivo="Reporte_Tienda_Vinilos.docx"
            documento.save(nombre_archivo)
            cursor.close()

            print(f"\n¡Reporte generado con éxito! Guardado como '{nombre_archivo}'.")
            funciones.accionExitosa()
            funciones.espereTecla()
            return True

        except Exception as e:
            print(f"\n[Error al generar el reporte]: {e}")
            funciones.accionNOExitosa()
            funciones.espereTecla()
            return False
    else:
        print("\nNo hay conexión con la base de datos.")
        funciones.espereTecla()
        return False